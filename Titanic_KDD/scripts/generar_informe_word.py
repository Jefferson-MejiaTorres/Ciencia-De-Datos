#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Script para generar un documento Word profesional del informe KDD del Titanic
Sigue normas APA con portada, contraportada, tabla de contenidos, etc.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def set_page_margins(doc, top=1, bottom=1, left=1, right=1):
    """Configura los márgenes de la página (en pulgadas)"""
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)

def add_page_break(doc):
    """Agrega un salto de página"""
    doc.add_page_break()

def style_heading(paragraph, level=1):
    """Estiliza un párrafo como encabezado APA"""
    if level == 1:
        paragraph.style = 'Heading 1'
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        paragraph.style = 'Heading 2'
    elif level == 3:
        paragraph.style = 'Heading 3'
    
    for run in paragraph.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

def add_footer_number(doc):
    """Agrega número de página en el pie de página"""
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = ""
    
    # Agregar número de página alineado a la derecha
    run = footer_para.add_run()
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    # Usar campo para número de página
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)
    
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# ============================================================================
# CREAR DOCUMENTO
# ============================================================================
doc = Document()

# Configurar márgenes APA (1 pulgada en todos lados)
set_page_margins(doc, top=1, bottom=1, left=1, right=1)

# Agregar número de página en pie de página
add_footer_number(doc)

# ============================================================================
# PORTADA APA
# ============================================================================
# Universidad
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("UNIVERSIDAD DE PAMPLONA")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run.bold = True

# Código y materia
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Código FGA-154 v.00\nFacultad de Ciencias\nMateria: Ciencia de Datos")
run.font.size = Pt(11)
run.font.name = 'Times New Roman'

# Espacios
for _ in range(3):
    doc.add_paragraph()

# Título del trabajo
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("LABORATORIO 1: PROCESO KDD\nANÁLISIS DEL TITANIC")
run.font.size = Pt(14)
run.font.name = 'Times New Roman'
run.bold = True

# Espacios
for _ in range(3):
    doc.add_paragraph()

# Autores
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Jefferson David Mejia Torres\nCC 1004924867\n\nDaniel Felipe Contreras Caballero\nCC 1005077931")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

# Espacios
for _ in range(3):
    doc.add_paragraph()

# Profesor y fecha
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Profesor: Abelardo Ferreira\n\n10 de Abril de 2026")
run.font.size = Pt(11)
run.font.name = 'Times New Roman'

add_page_break(doc)

# ============================================================================
# CONTRAPORTADA / RESUMEN EJECUTIVO
# ============================================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("RESUMEN EJECUTIVO")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run.bold = True

doc.add_paragraph()

# Contenido del resumen
resumen = """Este informe presenta un análisis integral del proceso KDD (Knowledge Discovery in Databases) aplicado al dataset histórico del RMS Titanic. El objetivo principal es identificar las variables demográficas, socioeconómicas y logísticas que fueron críticas en la probabilidad de supervivencia durante el naufragio del 15 de abril de 1912.

El análisis comprendió cinco fases: (1) Selección y descripción de datos (891 registros de pasajeros), (2) Preprocesamiento de datos incluyendo imputación estadística de valores faltantes, (3) Transformación de variables categóricas mediante codificación binaria y one-hot encoding, con la creación de cinco nuevas características mediante feature engineering, (4) Análisis exploratorio mediante estadísticas descriptivas y visualizaciones, y (5) Selección e implementación de un modelo de Regresión Logística.

Los resultados principales revelan que el género fue la variable más determinante (coeficiente = +2.494), con tasas de supervivencia de 74.20% para mujeres vs 18.89% para hombres (diferencia de 3.93 veces). La clase socioeconómica también ejerció un impacto significativo, mostrando supervivencia de 62.96% en primera clase versus 24.24% en tercera clase.

El modelo de Regresión Logística entrenado alcanzó una exactitud del 80.45% en los datos de prueba, con un ROC-AUC de 0.8569, indicando excelente capacidad discriminativa. El modelo identifica correctamente 4 de cada 5 casos y demuestra que la política de "mujeres y niños primero" fue implementada y fue efectiva durante la evacuación.

Cantidad de variables finales: 15 (5 nuevas características creadas)
División de datos: 80% entrenamiento (712 registros), 20% prueba (179 registros)
Desempeño final: 80.45% exactitud, F1-Score = 0.7287, ROC-AUC = 0.8569"""

p = doc.add_paragraph(resumen)
p_format = p.paragraph_format
p_format.line_spacing = 2.0
for run in p.runs:
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

add_page_break(doc)

# ============================================================================
# TABLA DE CONTENIDOS
# ============================================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("TABLA DE CONTENIDOS")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run.bold = True

doc.add_paragraph()

# Tabla de contenidos manual
toc_items = [
    ("1. Introducción", 0),
    ("2. Objetivo General", 0),
    ("   2.1 Objetivos Específicos", 1),
    ("3. Comprensión del Problema", 0),
    ("   3.1 Contexto del Dataset", 1),
    ("   3.2 Distribución del Dataset", 1),
    ("4. Descripción de Datos", 0),
    ("   4.1 Análisis General del Dataset", 1),
    ("   4.2 Análisis de Valores Nulos", 1),
    ("   4.3 Estadísticas Descriptivas", 1),
    ("   4.4 Validación de Datos", 1),
    ("5. Proceso KDD Detallado", 0),
    ("   5.1 Paso 1: Selección y Descripción de Datos", 1),
    ("   5.2 Paso 2: Preprocesamiento de Datos", 1),
    ("   5.3 Paso 3: Transformación y Feature Engineering", 1),
    ("   5.4 Paso 4: Análisis Exploratorio (EDA)", 1),
    ("   5.5 Paso 5: Selección y Modelado", 1),
    ("6. Visualizaciones", 0),
    ("7. Resultados y Descubrimientos", 0),
    ("   7.1 Evaluación del Modelo", 1),
    ("   7.2 Variables Críticas Identificadas", 1),
    ("   7.3 Análisis de Patrones de Supervivencia", 1),
    ("8. Conclusiones", 0),
    ("9. Referencias y Anexos", 0),
]

for item, indent in toc_items:
    p = doc.add_paragraph(item)
    p_format = p.paragraph_format
    if indent == 1:
        p_format.left_indent = Inches(0.5)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'

add_page_break(doc)

# ============================================================================
# 1. INTRODUCCIÓN
# ============================================================================
p = doc.add_paragraph()
style_heading(p, level=1)
p.clear()
run = p.add_run("1. INTRODUCCIÓN")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run.bold = True

intro_text = """El hundimiento del RMS Titanic el 15 de abril de 1912 es uno de los desastres marítimos más infames de la historia. Durante su viaje inaugural, el barco considerado "insumergible" chocó contra un iceberg y se hundió, causando la muerte de aproximadamente 1,502 de los 2,224 pasajeros y tripulantes a bordo.

Este laboratorio aplica el Proceso KDD (Knowledge Discovery in Databases) al dataset histórico del Titanic para realizar un análisis integral que permita identificar las variables críticas que influyeron en la probabilidad de supervivencia durante el naufragio.

Contexto del Desastre:
• Fecha: 15 de abril de 1912
• Lugar: Océano Atlántico Norte
• Causa: Colisión con iceberg
• Víctimas: 1,502 de 2,224 (67.5% tasa de mortalidad)
• Factor crítico: Solo 1,178 lugares en botes salvavidas para 2,224 personas

Durante este análisis, se examinarán las características demográficas, socioeconómicas y logísticas de los pasajeros para determinar qué factores incrementaban la probabilidad de supervivencia. El proceso KDD permitirá extraer conocimiento valioso de los datos históricos y desarrollar un modelo predictivo que pueda identificar patrones significativos en la supervivencia."""

p = doc.add_paragraph(intro_text)
p_format = p.paragraph_format
p_format.line_spacing = 2.0
for run in p.runs:
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# ============================================================================
# 2. OBJETIVO GENERAL
# ============================================================================
p = doc.add_paragraph()
style_heading(p, level=1)
p.clear()
run = p.add_run("2. OBJETIVO GENERAL")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run.bold = True

objetivo_text = """Aplicar el proceso KDD al dataset del Titanic con el fin de determinar las variables demográficas, socioeconómicas y logísticas que fueron críticas en la probabilidad de supervivencia durante el naufragio, estableciendo un modelo predictivo de regresión logística que pueda identificar patrones y relaciones en los datos históricos."""

p = doc.add_paragraph(objetivo_text)
p_format = p.paragraph_format
p_format.line_spacing = 2.0
for run in p.runs:
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Objetivos específicos
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("2.1 Objetivos Específicos")
run.font.size = Pt(11)
run.font.name = 'Times New Roman'
run.bold = True

objetivos = [
    "Cargar y describir exhaustivamente el dataset del Titanic",
    "Realizar un análisis exploratorio para identificar patrones y relaciones",
    "Limpiar y preprocesar los datos, manejando valores faltantes y variables irrelevantes",
    "Transformar variables categóricas y crear nuevas características significativas mediante feature engineering",
    "Entrenar un modelo de regresión logística para predicción de supervivencia",
    "Evaluar el desempeño del modelo, extraer conclusiones e identificar variables críticas"
]

for obj in objetivos:
    p = doc.add_paragraph(obj, style='List Bullet')
    p_format = p.paragraph_format
    p_format.line_spacing = 2.0
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

# ============================================================================
# 3. COMPRENSIÓN DEL PROBLEMA
# ============================================================================
p = doc.add_paragraph()
style_heading(p, level=1)
p.clear()
run = p.add_run("3. COMPRENSIÓN DEL PROBLEMA")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run.bold = True

# 3.1 Contexto
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("3.1 Contexto del Dataset")
run.font.size = Pt(11)
run.font.name = 'Times New Roman'
run.bold = True

contexto_text = """El dataset contiene información detallada de 891 pasajeros del Titanic, incluyendo variables demográficas, socioeconómicas y logísticas."""

p = doc.add_paragraph(contexto_text)
p_format = p.paragraph_format
p_format.line_spacing = 2.0
for run in p.runs:
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Variables
p = doc.add_paragraph("Variables Demográficas:")
for run in p.runs:
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

variables_demo = [
    "Age: Edad del pasajero en años",
    "Sex: Género del pasajero (masculino/femenino)",
    "SibSp: Número de hermanos/cónyuge a bordo",
    "Parch: Número de padres/hijos a bordo"
]
for var in variables_demo:
    p = doc.add_paragraph(var, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

p = doc.add_paragraph("Variables Socioeconómicas:")
for run in p.runs:
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

variables_socio = [
    "Pclass: Clase del billete (1ª, 2ª, 3ª clase)",
    "Fare: Tarifa pagada en libras esterlinas",
    "Cabin: Número de camarote",
    "Embarked: Puerto de embarque (Cherbourg, Queenstown, Southampton)"
]
for var in variables_socio:
    p = doc.add_paragraph(var, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

p = doc.add_paragraph("Variable Objetivo:")
for run in p.runs:
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

p = doc.add_paragraph("Survived: 1 = Sobrevivió, 0 = No sobrevivió", style='List Bullet')
for run in p.runs:
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

# 3.2 Distribución
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("3.2 Distribución del Dataset")
run.font.size = Pt(11)
run.font.name = 'Times New Roman'
run.bold = True

tabla_dist = [
    ("Total de registros", "891"),
    ("Total de variables", "12"),
    ("Tasa de supervivencia", "38.38% (342 sobrevivientes)"),
    ("Tasa de mortalidad", "61.62% (549 no sobrevivientes)")
]

table = doc.add_table(rows=len(tabla_dist)+1, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Métrica"
hdr_cells[1].text = "Valor"

for i, (metrica, valor) in enumerate(tabla_dist, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = metrica
    row_cells[1].text = valor

# Guardar página actual para continuar en siguiente
add_page_break(doc)

# ============================================================================
# 4. DESCRIPCIÓN DE DATOS
# ============================================================================
p = doc.add_paragraph()
style_heading(p, level=1)
p.clear()
run = p.add_run("4. DESCRIPCIÓN DE DATOS")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run.bold = True

p = doc.add_paragraph("4.1 Análisis General del Dataset")
for run in p.runs:
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

tabla_general = [
    ("Número de filas", "891"),
    ("Número de variables", "12"),
    ("Tipos de datos", "5 enteros, 2 flotantes, 5 objetos"),
    ("Memoria total", "83.7 KB")
]

table = doc.add_table(rows=len(tabla_general)+1, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Métrica"
hdr_cells[1].text = "Valor"

for i, (metrica, valor) in enumerate(tabla_general, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = metrica
    row_cells[1].text = valor

doc.add_paragraph()

p = doc.add_paragraph("4.2 Análisis de Valores Nulos")
for run in p.runs:
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

tabla_nulos = [
    ("Age", "177", "19.87%", "Imputar con mediana"),
    ("Cabin", "687", "77.10%", "Eliminar variable"),
    ("Embarked", "2", "0.22%", "Imputar con moda"),
    ("Resto", "0", "0%", "N/A")
]

table = doc.add_table(rows=len(tabla_nulos)+1, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Variable"
hdr_cells[1].text = "Nulos"
hdr_cells[2].text = "Porcentaje"
hdr_cells[3].text = "Estrategia"

for i, (var, nulos, porcentaje, estrategia) in enumerate(tabla_nulos, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = var
    row_cells[1].text = nulos
    row_cells[2].text = porcentaje
    row_cells[3].text = estrategia

doc.add_paragraph()
doc.add_paragraph("Acción: Se implementó imputación estadística conservadora para minimizar sesgo.")

add_page_break(doc)

# ============================================================================
# 5. PROCESO KDD DETALLADO
# ============================================================================
p = doc.add_paragraph()
style_heading(p, level=1)
p.clear()
run = p.add_run("5. PROCESO KDD DETALLADO")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run.bold = True

# Paso 1
p = doc.add_paragraph("5.1 Paso 1: Selección y Descripción de Datos")
for run in p.runs:
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

paso1_text = """Se seleccionaron las 12 variables originales del dataset:
• 10 variables relevantes para análisis predictivo
• 2 variables identificadores (PassengerId, Ticket) a eliminar posteriormente

Análisis de Calidad de Datos:
• Completitud: 91.8%
• Validez: 100% (tipos de datos correctos)
• Consistencia: 99.9% (sin registros completamente duplicados)
• Singularidad: 100% (cada pasajero es único)"""

p = doc.add_paragraph(paso1_text)
p_format = p.paragraph_format
p_format.line_spacing = 1.5
for run in p.runs:
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

doc.add_paragraph()

# Paso 2
p = doc.add_paragraph("5.2 Paso 2: Preprocesamiento de Datos")
for run in p.runs:
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

paso2_text = """Eliminación de Variables Irrelevantes:
• PassengerId: Solo ID, sin valor predictivo → ELIMINAR
• Name: Texto personal, no generalizable → ELIMINAR
• Ticket: Número único, sin patrones útiles → ELIMINAR
• Cabin: 77% nulos, información redundante → ELIMINAR

Resultado: Dataset reducido de 12 a 8 variables esenciales

Imputación de Valores Faltantes:

Age (19.87% nulos → 177 registros):
◦ Estrategia: Imputación con MEDIANA (28 años)
◦ Justificación: La mediana es más robusta ante outliers
◦ Impacto: Preserva la distribución sin sesgar

Embarked (0.22% nulos → 2 registros):
◦ Estrategia: Imputación con MODA (S = Southampton)
◦ Justificación: Puerto más frecuente
◦ Impacto: Negligible, <0.3% del dataset

Variable Derivada Creada:
FamilySize = SibSp + Parch + 1"""

p = doc.add_paragraph(paso2_text)
p_format = p.paragraph_format
p_format.line_spacing = 1.5
for run in p.runs:
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

doc.add_paragraph()

# Paso 3
p = doc.add_paragraph("5.3 Paso 3: Transformación y Feature Engineering")
for run in p.runs:
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

paso3_text = """Codificación de Variables Categóricas:

Variable Sex (Codificación Binaria):
female → 1 (presente) | male → 0 (ausente)
Correlación con supervivencia: +0.5434 ⭐ MUY FUERTE

Variable Embarked (One-Hot Encoding):
• Embarked_C: Puerto = Cherbourg (C) → 0 o 1
• Embarked_Q: Puerto = Queenstown (Q) → 0 o 1
• Embarked_S: Puerto = Southampton (S) → 0 o 1

Feature Engineering - Nuevas Características (5 creadas):

1) IsChild: 1 si Age < 15, sino 0
   - Frecuencia: 78 niños (8.8%)
   - Correlación: +0.1230
   - Justificación: Política "mujeres y niños primero"

2) IsAlone: 1 si FamilySize = 1, sino 0
   - Frecuencia: 537 pasajeros (60.3%)
   - Correlación: -0.2034
   - Justificación: Menor movilidad sin ayuda familiar

3) FamilyGroupSize: Categorización de familias
   - Grupo 1 (Solo): FamilySize = 1
   - Grupo 2 (Pequeño): FamilySize = 2-3
   - Grupo 3 (Mediano): FamilySize = 4-5
   - Grupo 4 (Grande): FamilySize ≥ 6

4) HasCabin: 1 si Cabin ≠ NULL, sino 0
   - Frecuencia: 204 pasajeros (22.9%)
   - Correlación: +0.3169
   - Justificación: Indicador de estatus e información

5) HigherClass: 1 si Pclass ≤ 2, sino 0
   - Frecuencia: 400 pasajeros (44.9%)
   - Correlación: +0.3223
   - Justificación: Simplifica a binaria, mejor acceso"""

p = doc.add_paragraph(paso3_text)
p_format = p.paragraph_format
p_format.line_spacing = 1.5
for run in p.runs:
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

add_page_break(doc)

# Paso 4
p = doc.add_paragraph("5.4 Paso 4: Análisis Exploratorio (EDA)")
for run in p.runs:
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

paso4_text = """Análisis Demográfico - SEXO:

Tasa de Supervivencia por Sexo:
• Mujeres: 233 (74.20%) de 314 total
• Hombres: 109 (18.89%) de 577 total
• Conclusión: Mujeres tuvieron 3.93 veces más probabilidad de sobrevivir
• Significancia estadística: χ² = 260.72, p-valor < 0.0001 ⭐ CRÍTICO

Análisis Socioeconómico - CLASE:

Tasa de Supervivencia por Clase:
• Primera: 136 (62.96%) de 216
• Segunda: 87 (47.28%) de 184
• Tercera: 119 (24.24%) de 491
• Conclusión: Primera clase 2.6x mejor que Tercera
• Significancia estadística: χ² = 102.89, p-valor < 0.0001 ⭐ CRÍTICO

Análisis de EDAD:
• No sobrevivieron: promedio 30.63 años (Des. Est. 14.17)
• Sobrevivieron: promedio 28.34 años (Des. Est. 14.95)
• Hallazgo: Niños < 10 años tenían > 60% supervivencia
• Significancia: t = -2.067, p-valor = 0.039 ⭐ SIGNIFICATIVO

Análisis de TARIFA (en Libras Esterlinas):
• No sobrevivieron: promedio £22.12
• Sobrevivieron: promedio £48.40
• Conclusión: Tarifa 2.19x mayor para sobrevivientes
• Significancia: t = 7.94, p-valor < 0.0001 ⭐ CRÍTICO

Patrones Observados:
1. Política de "mujeres y niños primero" fue implementada
2. Efecto de clase muy marcado en acceso a botes
3. Mayor tarifa correlacionada con mejor ubicación
4. Familias pequeñas más móviles que grandes grupos"""

p = doc.add_paragraph(paso4_text)
p_format = p.paragraph_format
p_format.line_spacing = 1.5
for run in p.runs:
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

doc.add_paragraph()

# Paso 5
p = doc.add_paragraph("5.5 Paso 5: Selección y Modelado")
for run in p.runs:
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

paso5_text = """División del Dataset (Stratified):
• Entrenamiento: 712 registros (79.9%)
• Prueba: 179 registros (20.1%)
• Estratificación: Mantenida para preservar distribución de clases
  ◦ Entrenamiento: 61.7% clase 0, 38.3% clase 1
  ◦ Prueba: 61.5% clase 0, 38.5% clase 1

Algoritmo Seleccionado: Regresión Logística

Justificación:
• Clasificación binaria (sobrevivió / no sobrevivió)
• Interpretabilidad clara de coeficientes
• Resultados probabilísticos naturales
• Desempeño robusto con este tipo de datos

Configuración del Modelo:
LogisticRegression(
    random_state=42,    # Reproducibilidad
    max_iter=1000,      # Iteraciones suficientes
    C=1.0              # Regularización L2 estándar
)

Convergencia: ✓ Exitosa en 4 iteraciones"""

p = doc.add_paragraph(paso5_text)
p_format = p.paragraph_format
p_format.line_spacing = 1.5
for run in p.runs:
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

add_page_break(doc)

# ============================================================================
# 6. VISUALIZACIONES
# ============================================================================
p = doc.add_paragraph()
style_heading(p, level=1)
p.clear()
run = p.add_run("6. VISUALIZACIONES")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run.bold = True

viz_intro = """Este proyecto genera 7 gráficos profesionales que facilitan la interpretación del proceso KDD y los resultados. Cada visualización comunica hallazgos específicos sobre la supervivencia en el Titanic."""

p = doc.add_paragraph(viz_intro)
p_format = p.paragraph_format
p_format.line_spacing = 2.0
for run in p.runs:
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Gráfico 1
p = doc.add_paragraph("Gráfico 1: Análisis de Supervivencia - Variables Demográficas y Socioeconómicas")
for run in p.runs:
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

grafico1_text = """Archivo: 01_supervivencia_basico.png

Contenido: 4 subgráficos mostrando:
• Supervivencia por Sexo: Comparativa de cantidad de sobrevivientes
• Tasa de Supervivencia por Sexo: Mujeres 74.2% vs Hombres 18.89%
• Supervivencia por Clase: Impacto socioeconómico
• Tasa de Supervivencia por Clase: 1ª clase 62.96%, 3ª clase 24.24%

Interpretación: El género y la clase fueron determinantes en la supervivencia."""

p = doc.add_paragraph(grafico1_text)
p_format = p.paragraph_format
p_format.line_spacing = 1.5
for run in p.runs:
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

# Espacio para imagen
p = doc.add_paragraph()
p = doc.add_paragraph("[ESPACIO RESERVADO PARA GRÁFICO 1]")
for run in p.runs:
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()

# Gráfico 2
p = doc.add_paragraph("Gráfico 2: Análisis de Distribuciones - Edad y Tarifa")
for run in p.runs:
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

grafico2_text = """Archivo: 02_distribuciones.png

Contenido: 2 histogramas superpuestos
• Distribución de Edad (0-80 años): Sobrevivientes vs No sobrevivientes
• Distribución de Tarifa (£0-£512): Concentración en diferentes rangos

Observaciones clave:
• Pico notable en edades 0-10 (niños con alta supervivencia)
• Sobrevivientes concentrados en tarifas mayores
• Tarifa promedio 2.19x mayor para sobrevivientes

Interpretación: La edad (especialmente niños) y la tarifa fueron factores de supervivencia."""

p = doc.add_paragraph(grafico2_text)
p_format = p.paragraph_format
p_format.line_spacing = 1.5
for run in p.runs:
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p = doc.add_paragraph("[ESPACIO RESERVADO PARA GRÁFICO 2]")
for run in p.runs:
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()

# Gráfico 3
p = doc.add_paragraph("Gráfico 3: Matriz de Correlación - Variables Numéricas")
for run in p.runs:
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

grafico3_text = """Archivo: 03_correlacion.png

Contenido: Heatmap 7x7 con correlaciones pairwise

Correlaciones clave con Survived:
• Pclass: -0.34 (Clase negativa = menor supervivencia)
• Age: -0.08 (Edad mayor = menor supervivencia, pero débil)
• Fare: +0.26 (Mayor tarifa = mayor supervivencia)
• FamilySize: +0.02 (Efecto casi nulo directo)

Color scheme:
🔴 Rojo oscuro: Correlación positiva fuerte
🔵 Azul oscuro: Correlación negativa fuerte
⚪ Beige: Correlación débil o nula

Interpretación: Multicolinealidad entre SibSp/Parch pero relaciones importantes con target."""

p = doc.add_paragraph(grafico3_text)
p_format = p.paragraph_format
p_format.line_spacing = 1.5
for run in p.runs:
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p = doc.add_paragraph("[ESPACIO RESERVADO PARA GRÁFICO 3]")
for run in p.runs:
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(128, 128, 128)

add_page_break(doc)

# Gráfico 4
p = doc.add_paragraph("Gráfico 4: Matriz de Confusión - Evaluación del Modelo")
for run in p.runs:
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

grafico4_text = """Archivo: 04_matriz_confusion.png

Contenido: Matriz 2x2 con escala de colores (azul oscuro = valores altos)

Resultados en Conjunto de Prueba (n=179):
• Verdaderos Negativos (TN): 97 ✓ Correctamente clasificados
• Falsos Positivos (FP): 13 ✗ Error tipo I
• Falsos Negativos (FN): 22 ✗ Error tipo II
• Verdaderos Positivos (TP): 47 ✓ Correctamente identificados

Métricas Derivadas:
• Exactitud: 80.45% = (97+47)/179
• Precisión: 78.33% = 47/(47+13)
• Recall (Sensibilidad): 68.12% = 47/(47+22)
• Especificidad: 88.18% = 97/(97+13)

Interpretación: El modelo clasifica correctamente 4 de 5 casos; mejor en identificar no-sobrevivientes."""

p = doc.add_paragraph(grafico4_text)
p_format = p.paragraph_format
p_format.line_spacing = 1.5
for run in p.runs:
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p = doc.add_paragraph("[ESPACIO RESERVADO PARA GRÁFICO 4]")
for run in p.runs:
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()

# Gráfico 5
p = doc.add_paragraph("Gráfico 5: Importancia de Variables - Coeficientes del Modelo")
for run in p.runs:
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

grafico5_text = """Archivo: 05_importancia_variables.png

Contenido: Gráfico de barras horizontal con coeficientes de Regresión Logística

Ranking de Importancia (por magnitud):
1. Sex: +2.494 ⭐⭐⭐ CRÍTICA - Impacto 12.16x en probabilidad
2. HasCabin: +0.904 ⭐⭐ Importante
3. IsChild: +0.842 ⭐⭐ Importante
4. IsAlone: -0.725 ⭐⭐ Reduce supervivencia
5. HigherClass: +0.643 ⭐ Relevante
6. FamilyGroupSize: -0.435 ⭐ Reduce supervivencia
7. Embarked_S: -0.425 ⭐ Reduce supervivencia
8. Pclass: -0.364 ⭐ Reduce supervivencia

Código de colores:
🟢 Verde: Coeficientes positivos (aumentan supervivencia)
🔴 Rojo: Coeficientes negativos (reducen supervivencia)

Conclusión: Sex domina completamente el modelo."""

p = doc.add_paragraph(grafico5_text)
p_format = p.paragraph_format
p_format.line_spacing = 1.5
for run in p.runs:
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p = doc.add_paragraph("[ESPACIO RESERVADO PARA GRÁFICO 5]")
for run in p.runs:
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()

# Gráfico 6
p = doc.add_paragraph("Gráfico 6: Curva ROC - Análisis del Modelo")
for run in p.runs:
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

grafico6_text = """Archivo: 06_curva_roc.png

Contenido: Curva ROC (Receiver Operating Characteristic)

Ejes:
• X (Horizontal): Tasa de Falsos Positivos (1 - Especificidad)
• Y (Vertical): Tasa de Verdaderos Positivos (Sensibilidad)

Características:
• Línea azul: Curva ROC del modelo logístico
• Línea punteada gris: Clasificador aleatorio (referencia = 0.50)
• AUC (Área bajo la curva): 0.8569 ⭐⭐⭐ EXCELENTE

Interpretación del AUC:
• 0.50: Clasificador completamente aleatorio
• 0.70-0.80: Discriminación aceptable
• 0.80-0.90: Discriminación excelente ← NUESTRO MODELO
• 0.90-1.00: Discriminación excepcional

Conclusión: AUC=0.8569 indica EXCELENTE capacidad predictiva."""

p = doc.add_paragraph(grafico6_text)
p_format = p.paragraph_format
p_format.line_spacing = 1.5
for run in p.runs:
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p = doc.add_paragraph("[ESPACIO RESERVADO PARA GRÁFICO 6]")
for run in p.runs:
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(128, 128, 128)

add_page_break(doc)

# Gráfico 7
p = doc.add_paragraph("Gráfico 7: Distribución de Probabilidades Predichas")
for run in p.runs:
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

grafico7_text = """Archivo: 07_distribucion_probabilidades.png

Contenido: Histograma superpuesto de probabilidades predichas

Ejes:
• X: Probabilidad Predicha de Supervivencia (0 a 1)
• Y: Frecuencia (cantidad de pasajeros)
• Colores: Rojo = No sobrevivieron, Verde = Sobrevivieron

Patrones Observados:

No sobrevivientes (rojo):
◦ Mayría concentrada en probabilidades bajas (0 a 0.3)
◦ Pico alto alrededor de 0.1
◦ Cola derecha (falsos negativos)

Sobrevivientes (verde):
◦ Mayoría concentrada en probabilidades altas (0.7 a 1.0)
◦ Pico importante alrededor de 0.85
◦ Cola izquierda (falsos positivos)

Línea punteada vertical: Umbral de decisión (probabilidad = 0.5)

Interpretación: La separación clara entre distribuciones indica que el modelo aprende patrones significativos y realiza predicciones bien calibradas."""

p = doc.add_paragraph(grafico7_text)
p_format = p.paragraph_format
p_format.line_spacing = 1.5
for run in p.runs:
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p = doc.add_paragraph("[ESPACIO RESERVADO PARA GRÁFICO 7]")
for run in p.runs:
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()
p = doc.add_paragraph()
p = doc.add_paragraph("Resumen de Gráficos:")
p = doc.add_paragraph("✅ 7 visualizaciones PNG generadas automáticamente")
p = doc.add_paragraph("✅ Resolución: Alto DPI para impresión profesional")
p = doc.add_paragraph("✅ Formato: Compatible con cualquier editor/presentador")
p = doc.add_paragraph("✅ Carpeta: Titanic_KDD/graficos/")

add_page_break(doc)

# ============================================================================
# 7. RESULTADOS Y DESCUBRIMIENTOS
# ============================================================================
p = doc.add_paragraph()
style_heading(p, level=1)
p.clear()
run = p.add_run("7. RESULTADOS Y DESCUBRIMIENTOS")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run.bold = True

p = doc.add_paragraph("7.1 Evaluación del Modelo")
for run in p.runs:
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

tabla_metricas = [
    ("Exactitud", "82.16%", "82.16%", "80.45%", "Modelo clasifica correctamente 4 de 5 casos"),
    ("Precisión", "-", "78.33%", "78.33%", "De 100 predichos como 'sobrevivió', 78 realmente sobrevivieron"),
    ("Recall (Sensibilidad)", "-", "68.12%", "68.12%", "El modelo detecta 68 de cada 100 sobrevivientes"),
    ("F1-Score", "-", "-", "0.7287", "Balance sólido entre precisión y recall"),
    ("ROC-AUC", "-", "-", "0.8569", "Excelente capacidad discriminativa")
]

table = doc.add_table(rows=len(tabla_metricas)+1, cols=5)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Métrica"
hdr_cells[1].text = "Entrenamiento"
hdr_cells[2].text = "Prueba"
hdr_cells[3].text = "Valor"
hdr_cells[4].text = "Interpretación"

for i, (metrica, entreno, prueba, valor, interp) in enumerate(tabla_metricas, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = metrica
    row_cells[1].text = entreno
    row_cells[2].text = prueba
    row_cells[3].text = valor
    row_cells[4].text = interp

doc.add_paragraph()

p = doc.add_paragraph("7.2 Variables Críticas Identificadas")
for run in p.runs:
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

ranking_vars = """Ranking de Importancia por Valor Absoluto del Coeficiente:

1. SEX (+2.4943) - ⭐⭐⭐ CRÍTICA
   • Ser mujer aumenta las 12.16 probabilidades de supervivencia (e^2.494)
   • Efecto dominante en el modelo
   • Refleja política de "mujeres primero" implementada durante evacuación

2. HASCABIN (+0.9036) - ⭐⭐ IMPORTANTE
   • Tener información de camarote aumenta 2.47 veces la probabilidad
   • Indicador de estatus y acceso a información

3. ISCHILD (+0.8421) - ⭐⭐ IMPORTANTE
   • Ser niño (< 15 años) aumenta 2.32 veces la probabilidad
   • Refleja política de "niños primero"

4. ISALONE (-0.7251) - ⭐⭐ IMPORTANTE
   • Viajar solo reduce 0.48 veces la probabilidad (48% menos)
   • Pasajeros solos menos móviles y sin apoyo familiar

5. HIGHERCLASS (+0.6426) - ⭐ RELEVANTE
   • Clase 1 o 2 aumenta 1.90 veces la probabilidad
   • Acceso preferente a botes, mejor información

6. FAMILYGROUPSIZE (-0.4348) - ⭐ RELEVANTE
   • Grupos grandes reducen probabilidad de supervivencia
   • Menor movilidad y coordinación

7. EMBARKED_S (-0.4250) - ⭐ RELEVANTE
   • Puerto Southampton muestra efecto negativo
   • Composición demográfica diferente por puerto

8. PCLASS (-0.3636) - ⭐ RELEVANTE
   • Clase de billete reduce probabilidad
   • Efecto capturado mejor por HigherClass"""

p = doc.add_paragraph(ranking_vars)
p_format = p.paragraph_format
p_format.line_spacing = 1.5
for run in p.runs:
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

doc.add_paragraph()

p = doc.add_paragraph("7.3 Análisis de Patrones de Supervivencia")
for run in p.runs:
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

patrones_text = """Perfil de Alto Riesgo:
• Hombre (Sex=0)
• Tercera clase (Pclass=3)
• Sin información de camarote
• Viajando solo
• Tarifa baja (< £10)
→ Probabilidad estimada de supervivencia: < 15%

Perfil de Bajo Riesgo:
• Mujer (Sex=1)
• Primera clase (Pclass=1)
• Con información de camarote
• Viajando con familia
• Tarifa alta (> £50)
→ Probabilidad estimada de supervivencia: > 90%

Principales Descubrimientos:

1. Género fue determinante: Variable más importante por gran margen

2. Clase socioeconómica fue crítica: Acceso diferenciado a recursos y botes

3. Edad y familia importaban: Niños y grupos pequeños mejor posicionados

4. Información de camarote indicaba estatus: Correlación con supervivencia

5. Puerto de embarque tuvo efecto: Diferentes composiciones demográficas por puerto

6. Política de evacuación observable: "Mujeres y niños primero" fue efectivamente implementada

7. Factor económico significativo: Mayor tarifa correlacionaba con mejor ubicación en el barco"""

p = doc.add_paragraph(patrones_text)
p_format = p.paragraph_format
p_format.line_spacing = 1.5
for run in p.runs:
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

add_page_break(doc)

# ============================================================================
# 8. CONCLUSIONES
# ============================================================================
p = doc.add_paragraph()
style_heading(p, level=1)
p.clear()
run = p.add_run("8. CONCLUSIONES")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run.bold = True

conclusiones_text = """El proceso KDD aplicado al dataset del Titanic ha permitido identificar con claridad las variables demográficas, socioeconómicas y logísticas que influyeron significativamente en la supervivencia durante el naufragio.

Hallazgos Principales:

Factor Demográfico Dominante:
El género fue absolutamente determinante (coeficiente = +2.494). La política de "mujeres y niños primero" se implementó efectivamente y fue documentalmente observable en los datos. Las mujeres tuvieron una probabilidad de supervivencia 3.93 veces mayor que los hombres (74.20% vs 18.89%).

Factor Socioeconómico Crítico:
La clase del billete dividió claramente las probabilidades de supervivencia. Pasajeros de 1ª clase tuvieron una tasa de supervivencia 2.6 veces mayor que los de 3ª clase (62.96% vs 24.24%). Esto refleja diferencias en acceso a botes salvavidas, información del desastre, y ayuda de la tripulación.

Factor Logístico Relevante:
El conocimiento del número de camarote indicaba mejor posicionamiento en el barco (coeficiente +0.904). El puerto de embarque mostró diferencias en composición demográfica, aunque con menor impacto individual.

Evaluación del Modelo Predictivo:

El modelo de Regresión Logística entrenado alcanzó:
• Exactitud del 80.45% en datos de prueba
• ROC-AUC de 0.8569 (excelente discriminación)
• F1-Score de 0.7287 (balance sólido entre precisión y recall)
• Precision del 78.33% (bajos falsos positivos)
• Recall del 68.12% (detecta mayoría de sobrevivientes)

Conclusión: El modelo es ROBUSTO y GENERALIZABLE a nuevos datos.

Implicaciones Históricas:

Los hallazgos corroboran el relato histórico:
✓ Política de "mujeres y niños primero" fue implementada con éxito
✓ Pasajeros de clase alta recibieron mejor trato y acceso a información
✓ La información y el acceso fueron literalmente factores de vida o muerte
✓ La tragedia no fue aleatoria; hubo patrones claros y medibles

Proceso KDD - Valor Agregado:

El análisis KDD no solo creó un modelo predictivo útil, sino que extrajo CONOCIMIENTO VALIOSO sobre los factores determinantes de supervivencia en este evento histórico crucial. A través de datos cuantitativos, pudimos validar narrativas históricas y cuantificar el impacto de variables sociales en resultados de vida o muerte.

Recomendaciones para Futuras Investigaciones:

1. Incorporar datos de tripulación vs. pasajeros para análisis comparativo
2. Investigar variables de ubicación en el barco (cubierta, proximidad a botes)
3. Estudiar correlaciones temporales (quién fue evacuado primero)
4. Comparar con datos de otros naufragios contemporáneos
5. Aplicar modelos más complejos (Random Forest, XGBoost) para benchmarking

Reflexión Final:

Este análisis demuestra cómo el DATA MINING puede revelar verdades históricas cuantificables sobre eventos del pasado. Las variables categóricas simples pero significativas (sexo, clase) muestran cómo la estructura social de 1912 impactó literalmente en probabilidades de vida o muerte durante la tragedia del Titanic. El proceso KDD proporciona una metodología rigurosa para extraer conocimiento de datos históricos y crear narrativas basadas en evidencia."""

p = doc.add_paragraph(conclusiones_text)
p_format = p.paragraph_format
p_format.line_spacing = 2.0
for run in p.runs:
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

add_page_break(doc)

# ============================================================================
# 9. REFERENCIAS Y ANEXOS
# ============================================================================
p = doc.add_paragraph()
style_heading(p, level=1)
p.clear()
run = p.add_run("9. REFERENCIAS Y ANEXOS")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run.bold = True

p = doc.add_paragraph("Referencias:")
for run in p.runs:
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

referencias = [
    "Kaggle Titanic Dataset. (2024). Recuperado de https://www.kaggle.com/c/titanic",
    "Encyclopedia Britannica. (2024). RMS Titanic. Recuperado de https://www.britannica.com/topic/Titanic",
    "U.S. Senate Inquiry into the Loss of the RMS Titanic. (1912). Informe histórico del Congreso de EE.UU.",
    "Witten, I. H., Frank, E., & Hall, M. A. (2011). Data Mining: Practical Machine Learning Tools and Techniques (3rd ed.). Morgan Kaufmann.",
    "Scikit-learn Development Team. (2024). Logistic Regression Documentation. Recuperado de https://scikit-learn.org/",
    "McKinney, W. (2012). Python for Data Analysis. O'Reilly Media.",
]

for ref in referencias:
    p = doc.add_paragraph(ref, style='List Bullet')
    p_format = p.paragraph_format
    p_format.line_spacing = 1.5
    p_format.left_indent = Inches(0.5)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'

doc.add_paragraph()

p = doc.add_paragraph("Anexo A: Estructura de Archivos del Proyecto")
for run in p.runs:
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

estructura_text = """Titanic_KDD/
├── INFORME_FINAL.md                        # Informe técnico completo
├── README.md                               # Guía de ejecución y proyecto
├── scripts/
│   ├── paso1_carga.py                     # Cargue y descripción inicial
│   ├── paso2_exploracion.py               # Análisis exploratorio (EDA)
│   ├── paso3_limpieza.py                  # Preprocesamiento de datos
│   ├── paso4_transformacion.py            # Feature engineering
│   └── paso5_modelado.py                  # Entrenamiento del modelo
├── datos/
│   ├── train.csv                          # Dataset original (891 registros)
│   ├── df_original.csv                    # Datos cargados
│   ├── df_limpio.csv                      # Después de limpieza
│   └── df_transformado.csv                # Transformado (listo para ML)
└── graficos/
    ├── 01_supervivencia_basico.png
    ├── 02_distribuciones.png
    ├── 03_correlacion.png
    ├── 04_matriz_confusion.png
    ├── 05_importancia_variables.png
    ├── 06_curva_roc.png
    └── 07_distribucion_probabilidades.png"""

p = doc.add_paragraph(estructura_text)
p_format = p.paragraph_format
p_format.line_spacing = 1.2
for run in p.runs:
    run.font.size = Pt(10)
    run.font.name = 'Courier New'
    run.font.color.rgb = RGBColor(64, 64, 64)

doc.add_paragraph()

p = doc.add_paragraph("Anexo B: Configuración Técnica del Proyecto")
for run in p.runs:
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

config_text = """Lenguaje de Programación: Python 3.13.3

Librerías Principales:
• pandas: Análisis y manipulación de datos
• numpy: Computación numérica y arrays
• scikit-learn: Machine Learning (Logistic Regression, train_test_split)
• matplotlib: Visualización estática (backend Agg para entornos headless)
• seaborn: Visualización estadística avanzada
• scipy: Análisis estadístico (chi2_contingency, ttest_ind)

Entorno Virtual:
• Ubicación: C:/Repositorios/Ciencia De Datos/.venv/
• Tipo: Python venv (virtual environment)
• Activación: source .venv/Scripts/activate

Método de División de Datos:
• Estrategia: Stratified Train-Test Split
• Proporción: 80% entrenamiento, 20% prueba
• Estratificación: Mantenida para preservar distribución de variable objetivo

Modelo de Machine Learning:
Algoritmo: Logistic Regression (sklearn)
Parámetros:
  - solver: lbfgs (por defecto)
  - max_iter: 1000 (iteraciones máximas)
  - C: 1.0 (inverso de fuerza de regularización)
  - random_state: 42 (reproducibilidad)
  - Regularización: L2

Métricas de Evaluación:
• Exactitud (Accuracy): Proporción correcta
• Precisión (Precision): De predicciones positivas, cuántas correctas
• Recall (Sensibilidad): De casos positivos reales, cuántos se detectan
• F1-Score: Media armónica de precisión y recall
• ROC-AUC: Área bajo la curva ROC

Especificaciones de Visualizaciones:
• Formato: PNG
• Resolución: Alto DPI (300+) para impresión
• Tamaño: Optimizado para reportes
• Biblioteca: Matplotlib y Seaborn"""

p = doc.add_paragraph(config_text)
p_format = p.paragraph_format
p_format.line_spacing = 1.3
for run in p.runs:
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'

# Guardar documento
output_path = r"c:\Repositorios\Ciencia De Datos\Titanic_KDD\INFORME_TITANIC_FINAL.docx"
doc.save(output_path)

print(f"✅ Documento Word generado exitosamente:")
print(f"   Ubicación: {output_path}")
print(f"\n📋 Contenido incluido:")
print(f"   ✓ Portada APA profesional")
print(f"   ✓ Contraportada / Resumen ejecutivo")
print(f"   ✓ Tabla de contenidos")
print(f"   ✓ 8 secciones principales")
print(f"   ✓ Espacios reservados para 7 gráficos")
print(f"   ✓ Tablas de datos y resultados")
print(f"   ✓ Referencias bibliográficas")
print(f"   ✓ Anexos técnicos")
print(f"   ✓ Formato profesional con márgenes APA")
print(f"   ✓ Numeración de páginas")
print(f"   ✓ Fuente Times New Roman 12pt")
print(f"   ✓ Espaciado doble")
